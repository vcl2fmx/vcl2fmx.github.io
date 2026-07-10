from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\New Delphi Projects\VCL2FMXConverterV5")
OUT = ROOT / "docs" / "VCL2FMXConverter_v5_0_Conversion_Contracts_Guide.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 35, 55)
MUTED = RGBColor(95, 105, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, col_widths=None):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    if col_widths:
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in col_widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx < len(col_widths):
                    tc_pr = cell._tc.get_or_add_tcPr()
                    tc_w = tc_pr.find(qn("w:tcW"))
                    if tc_w is None:
                        tc_w = OxmlElement("w:tcW")
                        tc_pr.append(tc_w)
                    tc_w.set(qn("w:w"), str(col_widths[idx]))
                    tc_w.set(qn("w:type"), "dxa")


def set_borders(table, color="B7C5D8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def add_static_toc(doc):
    items = [
        "1. Purpose of This Guide",
        "2. What a Contract Is",
        "3. Why v5.0 Uses Contracts",
        "4. How the Contract Runner Works",
        "5. Contract Folder Structure",
        "6. Anatomy of an Expectation File",
        "7. What Contracts Prove",
        "8. What Contracts Do Not Prove",
        "9. How Contracts Interact With Mapping Packs",
        "10. How Contracts Interact With Reports",
        "11. How to Add a New Contract",
        "12. Good Contract Design",
        "13. Weak Contracts and How to Repair Them",
        "14. Public Source Distribution Notes",
        "15. Recommended Review Instruction for Outside AI Reviewers",
        "16. Current v5.0 Contract Baseline",
        "17. Summary",
    ]
    for item in items:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.18)
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(item)
        run.font.color.rgb = DARK_BLUE


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor(25, 35, 45)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(4)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.line_spacing = 1.0

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = INK
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "VCL2FMXConverter v5.0 Vanguard - Conversion Contracts Guide"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = ""
    add_page_number(footer)
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("VCL2FMXConverter")
    run.font.name = "Calibri"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Conversion Contracts Guide")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Version 5.0 Vanguard - detailed engineering and operator reference")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    set_table_width(meta, col_widths=[2500, 6860])
    set_cell_margins(meta)
    set_borders(meta)
    rows = [
        ("Audience", "Developers, maintainers, beta testers, and reviewers who need to understand or extend the v5.0 contract system."),
        ("Revision date", "June 28, 2026"),
        ("Project folder", r"C:\New Delphi Projects\VCL2FMXConverterV5"),
        ("Contract root", r"C:\New Delphi Projects\VCL2FMXConverterV5\contracts"),
    ]
    for row, (left, right) in zip(meta.rows, rows):
        row.cells[0].text = left
        row.cells[1].text = right
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            for run in row.cells[0].paragraphs[0].runs:
                run.font.bold = True

    doc.add_paragraph()


def h(doc, level, text):
    return doc.add_heading(text, level=level)


def p(doc, text="", style=None):
    return doc.add_paragraph(text, style=style)


def bullet(doc, text, level=0):
    para = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    para.paragraph_format.space_after = Pt(4)
    para.add_run(text)
    return para


def number(doc, text):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_after = Pt(4)
    para.add_run(text)
    return para


def code(doc, text):
    for line in text.strip("\n").splitlines():
        para = doc.add_paragraph(style="Code Block")
        para.add_run(line)


def callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, col_widths=[9360])
    set_cell_margins(table, top=100, bottom=100, start=160, end=160)
    set_borders(table, "A9BCD4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(3)
    r = para.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    para.add_run(" " + text)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, col_widths=widths)
    set_cell_margins(table)
    set_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = DARK_BLUE
                run.font.size = Pt(10)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    h(doc, 1, "Table of Contents")
    add_static_toc(doc)
    doc.add_page_break()

    h(doc, 1, "1. Purpose of This Guide")
    p(doc, "This guide explains the v5.0 Vanguard conversion-contract system in practical terms: what contracts are, what they contain, how the runner uses them, what they prove, what they do not prove, and how new contracts should be written when the converter learns a new rule.")
    p(doc, "The contract system exists because a source-to-source converter can look correct on one project and still fail in a different code shape. A contract captures the shape of a known problem and turns it into an executable rule. Once a problem has a contract, the converter either keeps satisfying that rule or the contract runner fails.")
    callout(doc, "Plain-language definition:", "A conversion contract is a small source input case plus a written expectation file. The input case creates a specific conversion problem. The expectation file states exactly what the converted output and report must contain, and what must not appear.")

    h(doc, 1, "2. What a Contract Is")
    p(doc, "A VCL2FMX conversion contract has two parts.")
    number(doc, "A fixture: one or more source files such as .pas, .dfm, .dpr, .inc, or project companion files. The fixture deliberately contains one behavior the converter must handle.")
    number(doc, "An expectation: a matching .expected.json file that records the expected conversion status, generated-code patterns, forbidden patterns, report messages, and uses-clause results.")
    p(doc, "The fixture is the test input. The expectation is the rule. Together, they become an executable contract.")
    add_table(
        doc,
        ["Part", "Example", "Purpose"],
        [
            ("Fixture", r"contracts\colors\01_class_should_not_add_uiconsts.pas", "Provides a small Pascal unit containing a class declaration and non-color words like class and declare."),
            ("Expectation", r"contracts\colors\01_class_should_not_add_uiconsts.expected.json", "Forbids System.UIConsts and System.UITypes because the source does not use a real FMX cla color constant."),
            ("Generated output", r"tests\conversion_contract_output\...\output", "Temporary output created by the runner so the output can be checked and discarded."),
            ("Report output", "VCL_to_FMX_Conversion_Report.txt/html", "The conversion report checked for required manual-review items, warnings, errors, or forbidden false positives."),
        ],
        [1700, 3300, 4360],
    )

    h(doc, 2, "What a Contract Is Not")
    bullet(doc, "It is not runtime data loaded by the normal converter UI.")
    bullet(doc, "It is not a mapping pack.")
    bullet(doc, "It is not a sample project meant for users to convert by hand.")
    bullet(doc, "It is not compared against every user project during a normal conversion.")
    bullet(doc, "It is not a promise that the whole converted project is finished; it proves one rule or one defined cluster of rules.")

    h(doc, 1, "3. Why v5.0 Uses Contracts")
    p(doc, "The converter performs many different tasks: DFM parsing, FMX generation, Pascal rewriting, uses-clause cleanup, include-file analysis, Windows-message handling, reporting, mapping-pack detection, dry-run output, and more. A small change in one area can accidentally break a distant area.")
    p(doc, "Contracts reduce that risk by turning each important behavior into a repeatable check. They also make outside reviews more useful. When a reviewer finds a bug, the next question is no longer just whether the source code is wrong. The next question is whether an existing contract should have caught it.")
    bullet(doc, "If no contract exists, add one before trusting the fix.")
    bullet(doc, "If a contract exists but still passed, strengthen the contract.")
    bullet(doc, "If the contract fails after a change, decide whether the converter changed incorrectly or the expected behavior intentionally changed.")

    h(doc, 1, "4. How the Contract Runner Works")
    p(doc, "The contract runner is an internal test tool. It scans the contracts folder, finds fixtures with matching expectations, runs the converter against each fixture, and compares generated files and reports against the expectation JSON.")
    p(doc, "The runner does not use a special database index. It discovers contracts directly from the folder tree. This is fast enough because each fixture is small and the directory structure is simple.")
    add_table(
        doc,
        ["Step", "What Happens", "Why It Matters"],
        [
            ("1", "Discover expectation files under contracts.", "Each expectation identifies the input fixture to run."),
            ("2", "Copy the fixture into an isolated temporary source folder.", "The original contract fixture remains unchanged."),
            ("3", "Run the converter through RunConversionEngine.exe.", "The same core engine is exercised without clicking the UI."),
            ("4", "Read generated Pascal, FMX, DPR, and reports.", "The runner checks behavior, not just process completion."),
            ("5", "Apply expected and forbidden regex patterns.", "A contract can require correct output and forbid dangerous leftovers."),
            ("6", "Print pass/fail summary.", "Any unexpected behavior stops the release gate."),
        ],
        [900, 3900, 4560],
    )

    h(doc, 2, "Normal Conversion Versus Contract Execution")
    p(doc, "A normal user conversion does not compare the user's project to every contract. The converter applies its built-in rules, mapping packs, and analysis logic directly to the user's source. The contracts are only used when the internal contract runner is executed.")
    p(doc, "In other words, if there are 19 user projects and 150 contracts, the converter does not run 19 times 150 comparisons. The contracts are only test cases for the converter itself.")

    h(doc, 1, "5. Contract Folder Structure")
    p(doc, "The v5.0 contracts are organized by behavior family so a maintainer can find the right place for a new fixture quickly.")
    add_table(
        doc,
        ["Folder", "Purpose"],
        [
            ("colors", "Color constants, UI constants, modal results, and false positives such as class and clarity."),
            ("comments_and_boundaries", "Comments, strings, disabled blocks, routine boundaries, initialization/finalization boundaries, and false-positive guards."),
            ("components_and_events", "Component declarations, lifecycle events, media events, image/property assignments, shape conversion, and event wiring."),
            ("dfm_fidelity", "Checks for missing DFM/FMXL objects, properties, and events after generation."),
            ("dfm_pairs", "Paired Pascal/DFM fixtures for FMX form generation."),
            ("graphics", "Canvas/GDI conversion and visual-review reporting."),
            ("include_analysis", "Pascal include directives, missing/outside/nested/recursive includes, and UTF-8 include files."),
            ("project_integration", "Whole mini-project fixtures that exercise multiple systems together."),
            ("uses_clause", "Unit cleanup, Windows/VCL unit preservation/removal, and conditional uses behavior."),
            ("winapi_messages", "Windows messaging APIs, message constants, WndProc, message declarations, system commands, and false positives."),
        ],
        [2600, 6760],
    )

    h(doc, 1, "6. Anatomy of an Expectation File")
    p(doc, "Most expectations use the same fields. Fields can be omitted when they are not relevant to a fixture.")
    add_table(
        doc,
        ["Field", "Meaning"],
        [
            ("input_file", "The fixture path relative to the project root."),
            ("expected_status", "The conversion status expected by the runner, usually converted."),
            ("expected_conversions", "Conversion categories expected to be reported."),
            ("expected_manual_reviews", "Manual-review categories that must be present."),
            ("expected_units_added", "Units that must appear in the generated uses clause."),
            ("expected_units_absent", "Units that must not appear in the generated uses clause."),
            ("expected_output_patterns", "Regex patterns that must appear in generated output files."),
            ("forbidden_output_patterns", "Regex patterns that must not appear in generated output files."),
            ("expected_report_patterns", "Regex patterns that must appear in the text or HTML report."),
            ("forbidden_report_patterns", "Regex patterns that must not appear in reports."),
            ("copy_case_directory", "Optional project-level mode that copies an entire fixture directory before conversion."),
            ("skip_generated_compile", "Optional flag for fixtures that intentionally preserve missing or manual-review dependencies."),
        ],
        [3000, 6360],
    )

    h(doc, 2, "Example: A Simple Uses-Clause Guard")
    p(doc, "This contract proves that ordinary words beginning with cla do not inject FMX UI color units.")
    code(doc, r'''
{
  "input_file": "contracts\\colors\\01_class_should_not_add_uiconsts.pas",
  "expected_status": "converted",
  "expected_output_patterns": ["unit ContractClassShouldNotAddUIConsts"],
  "forbidden_output_patterns": [
    "System\\.UIConsts",
    "System\\.UITypes"
  ],
  "expected_report_patterns": [],
  "forbidden_report_patterns": []
}
''')
    p(doc, "This expectation is deliberately small. It does not care about every line of generated Pascal. It cares only that the unit still converts and that two false-positive units are absent.")

    h(doc, 2, "Example: A Positive Color Contract")
    p(doc, "A good negative contract should usually be paired with a positive contract. The negative contract says class should not trigger color units. The positive contract says a real FMX color constant should trigger them.")
    code(doc, r'''
procedure TContractRealClaColorShouldAddUIConsts.Run;
var
  ColorValue: TAlphaColor;
begin
  ColorValue := claBlack;
  Writeln(Integer(ColorValue));
end;
''')
    p(doc, "The expectation for this fixture requires both System.UIConsts and System.UITypes. Together, the positive and negative contracts prevent both over-matching and under-matching.")

    h(doc, 1, "7. What Contracts Prove")
    bullet(doc, "The converter recognizes a known input shape.")
    bullet(doc, "The converter produces required generated-code patterns.")
    bullet(doc, "The converter avoids known dangerous or noisy output patterns.")
    bullet(doc, "The report contains required manual-review, warning, blocking, or informational messages.")
    bullet(doc, "The generated uses clause includes needed units and excludes unused or false-positive units.")
    bullet(doc, "A past bug stays fixed after future changes.")

    h(doc, 1, "8. What Contracts Do Not Prove")
    bullet(doc, "They do not prove every real project will compile without manual review.")
    bullet(doc, "They do not replace beta testing against real applications.")
    bullet(doc, "They do not make unsupported VCL or Windows behavior automatically safe.")
    bullet(doc, "They do not prove visual correctness beyond the defined output/report patterns.")
    bullet(doc, "They do not prove third-party component behavior unless a mapping rule and fixture explicitly cover that behavior.")
    bullet(doc, "They do not run during normal user conversion unless the internal runner is invoked.")

    h(doc, 1, "9. How Contracts Interact With Mapping Packs")
    p(doc, "Mapping packs describe third-party components and conversion actions. Contracts verify converter behavior. They are related, but not the same.")
    add_table(
        doc,
        ["Mapping Pack", "Contract"],
        [
            ("Contains JSON rules for known third-party component classes.", "Contains fixtures and expectations that test converter behavior."),
            ("Used during normal conversion.", "Used by the internal test runner."),
            ("Can identify, preserve, substitute, or review a component.", "Can prove that mapping-pack behavior appears in generated output or reports."),
            ("Ships in the mapping_packs folder.", "Ships in the contracts folder as test evidence."),
        ],
        [4680, 4680],
    )
    p(doc, "When a mapping rule changes, a contract should be added or updated. The contract should prove the expected generated component, preserved property, manual-review message, or forbidden false positive.")

    h(doc, 1, "10. How Contracts Interact With Reports")
    p(doc, "The report is part of the converter's output. For migration work, a truthful report can be as important as generated code. Contracts can therefore assert report behavior.")
    bullet(doc, "A real unsupported item should appear in the report.")
    bullet(doc, "A false positive should not appear in the report.")
    bullet(doc, "A blocking item should be marked as blocking when generated output cannot be trusted.")
    bullet(doc, "A visual-review item should still report visual review even when a safe code conversion is possible.")
    bullet(doc, "Informational dry-run notices should not inflate actionable issue totals.")
    p(doc, "This distinction helped v5.0 separate warnings, manual-review items, blockers, informational notices, and dry-run notices.")

    h(doc, 1, "11. How to Add a New Contract")
    p(doc, "Add a new contract when a behavior is important enough that losing it would damage conversion quality, compile readiness, report honesty, or user trust.")
    number(doc, "Choose the correct behavior folder under contracts.")
    number(doc, "Create the smallest fixture that demonstrates the problem.")
    number(doc, "Create a matching .expected.json file.")
    number(doc, "Use expected patterns for behavior that must appear.")
    number(doc, "Use forbidden patterns for behavior that must never appear.")
    number(doc, "Run the focused contract folder first.")
    number(doc, "Run the full contract suite before trusting the change.")
    number(doc, "If the converter change touches shared logic, run regression guards too.")

    h(doc, 2, "Example Workflow for a New Bug")
    p(doc, "Suppose a converted project reveals that TShape.Shape = stCircle becomes TEllipse correctly, but the report still asks for manual review. The proper v5.0 workflow is:")
    number(doc, "Create a fixture with a TShape whose DFM Shape property is stCircle.")
    number(doc, "Add Pascal code that assigns shpLED.Shape := stCircle.")
    number(doc, "Expect TEllipse in the generated Pascal and FMX.")
    number(doc, "Forbid Shape = stCircle in generated FMX.")
    number(doc, "Forbid FMX manual review for shpLED.Shape.")
    number(doc, "Fix the converter.")
    number(doc, "Run contracts and regression guards.")

    h(doc, 1, "12. Good Contract Design")
    p(doc, "Good contracts are precise. They do not try to check everything. They check the behavior that matters.")
    add_table(
        doc,
        ["Good Practice", "Reason"],
        [
            ("Keep fixtures small.", "Small fixtures isolate the behavior and make failures easy to understand."),
            ("Use one main rule per contract.", "A contract that checks too much can fail for the wrong reason."),
            ("Include forbidden patterns.", "Many converter bugs are leftovers, not missing additions."),
            ("Pair positive and negative tests.", "This prevents over-matching and under-matching."),
            ("Assert report behavior.", "A clean compile is not enough if the report hides migration risk."),
            ("Use project-level fixtures when needed.", "Some behavior only appears when multiple units, includes, or DPR files interact."),
        ],
        [3000, 6360],
    )

    h(doc, 1, "13. Weak Contracts and How to Repair Them")
    p(doc, "A weak contract is a contract that passes even when the converter is wrong. The v5.0 color issue is a good example. The contract name said class should not add UI constants, but the expectation did not forbid System.UIConsts or System.UITypes. The runner passed because the expectation was too vague.")
    callout(doc, "Rule:", "A contract name is not a test. Only the expectation file is a test. If the expectation does not require or forbid the right pattern, the runner cannot catch the bug.")
    p(doc, "To repair a weak contract, update the expectation so it would have failed before the fix. Then run the old behavior mentally or with the old executable if available. A repaired contract should prove that the bug is genuinely guarded.")

    h(doc, 1, "14. Public Source Distribution Notes")
    p(doc, "The public source-distribution zip includes the contracts folder because contracts document the expected converter behavior. The public source zip intentionally excludes internal tools and generated outputs. That means the contracts are present as authoritative fixtures, but the runner scripts and RunConversionEngine.exe are not included in that compile-only zip.")
    p(doc, "For internal development, the runner remains in the tools folder of the working repository. Before trusting contract results after engine changes, rebuild both the main converter executable and the contract runner executable.")

    h(doc, 1, "15. Recommended Review Instruction for Outside AI Reviewers")
    p(doc, "When sending the project to another code reviewer, include the contracts folder and explain how to interpret it. A useful instruction is:")
    code(doc, r'''
Review the converter source together with the contracts folder.
Each contract is a fixture plus an .expected.json file describing required
output, forbidden output, report expectations, and uses-clause expectations.
If you find a bug, check whether a contract already covers it.
If a contract exists but would not fail, report that the contract is weak.
If no contract exists, propose a new contract fixture and expected JSON.
''')

    h(doc, 1, "16. Current v5.0 Contract Baseline")
    add_table(
        doc,
        ["Metric", "Current Value"],
        [
            ("Contract expectations", "189"),
            ("Mapping packs", "22"),
            ("Last verified full contract run", "189 pass, 0 fail"),
            ("Regression guard status", "41 pass, 0 blockers, 0 warnings"),
            ("Normal conversion dependency", "Contracts are not loaded during normal user conversion."),
        ],
        [3200, 6160],
    )

    h(doc, 1, "17. Summary")
    p(doc, "The v5.0 contract system is the converter's memory. It records known migration problems as executable examples, checks generated output and reports, and prevents the same class of error from silently returning. The contracts do not remove the need for real-project testing, but they make real-project discoveries permanent. Every important converter rule should eventually have a fixture, an expectation, and a regression guard.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
