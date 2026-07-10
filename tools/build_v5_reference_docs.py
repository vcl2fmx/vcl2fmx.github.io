from __future__ import annotations

import html
import json
import re
import subprocess
import zipfile
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
MAPPER = ROOT / "Converter.Mapper.Component.pas"
MAPPING_PACKS = ROOT / "mapping_packs"
CONTRACTS = ROOT / "contracts"
TODAY = date.today().strftime("%B %d, %Y")


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def pas_string(line: str) -> str:
    m = re.search(r":=\s*'((?:''|[^'])*)'", line)
    return m.group(1).replace("''", "'") if m else ""


def pas_bool(line: str) -> bool:
    return "True" in line


def pas_int(line: str) -> int:
    m = re.search(r":=\s*(\d+)", line)
    return int(m.group(1)) if m else 0


def parse_builtin_mappings() -> list[dict]:
    lines = read_text(MAPPER).splitlines()
    mappings: list[dict] = []
    current: dict | None = None
    prop: dict | None = None
    event: dict | None = None

    for raw in lines:
        line = raw.strip()
        if line.startswith("Mapping := TComponentMapping.Create"):
            current = {
                "source": "built-in",
                "vcl_class": "",
                "fmx_class": "",
                "mapping_type": "",
                "confidence": 0,
                "notes": "",
                "properties": [],
                "events": [],
            }
            prop = None
            event = None
            continue
        if current is None:
            continue
        if line.startswith("Mapping.VCLClassName"):
            current["vcl_class"] = pas_string(line)
        elif line.startswith("Mapping.FMXClassName"):
            current["fmx_class"] = pas_string(line)
        elif line.startswith("Mapping.MappingType"):
            current["mapping_type"] = pas_string(line)
        elif line.startswith("Mapping.Confidence"):
            current["confidence"] = pas_int(line)
        elif line.startswith("Mapping.Notes"):
            current["notes"] = pas_string(line)
        elif line.startswith("PropMap.VCLProp"):
            prop = {"vcl": pas_string(line), "fmx": "", "transform": False, "transformer": ""}
        elif line.startswith("PropMap.FMXProp") and prop is not None:
            prop["fmx"] = pas_string(line)
        elif line.startswith("PropMap.NeedsTransformation") and prop is not None:
            prop["transform"] = pas_bool(line)
        elif line.startswith("PropMap.TransformerFunc") and prop is not None:
            prop["transformer"] = pas_string(line)
        elif line.startswith("Mapping.PropertyMaps.Add") and prop is not None:
            current["properties"].append(prop)
            prop = None
        elif line.startswith("EventMap.VCLEvent"):
            event = {"vcl": pas_string(line), "fmx": "", "signature_match": True}
        elif line.startswith("EventMap.FMXEvent") and event is not None:
            event["fmx"] = pas_string(line)
        elif line.startswith("EventMap.SignatureMatch") and event is not None:
            event["signature_match"] = pas_bool(line)
        elif line.startswith("Mapping.EventMaps.Add") and event is not None:
            current["events"].append(event)
            event = None
        elif line.startswith("FMappingDatabase.Add(Mapping"):
            if current.get("vcl_class"):
                mappings.append(current)
            current = None
            prop = None
            event = None
    return mappings


def parse_mapping_packs() -> list[dict]:
    rows: list[dict] = []
    for path in sorted(MAPPING_PACKS.glob("*_MappingPack_v4_1.json")):
        data = json.loads(read_text(path))
        pack = data.get("display_name") or data.get("name") or path.stem
        vendor = data.get("vendor", "")
        for rule in data.get("rules", []):
            rows.append(
                {
                    "source": "mapping pack",
                    "pack": pack,
                    "vendor": vendor or rule.get("vendor", ""),
                    "vcl_class": rule.get("vcl_class", ""),
                    "fmx_class": rule.get("fmx_class", ""),
                    "mapping_type": rule.get("mapping_type", rule.get("action", "")),
                    "action": rule.get("action", ""),
                    "confidence": rule.get("confidence", 0),
                    "notes": rule.get("notes", ""),
                    "manual_review_reason": rule.get("manual_review_reason", ""),
                    "properties": [
                        {
                            "vcl": p.get("vcl", ""),
                            "fmx": p.get("fmx", ""),
                            "transform": p.get("needs_transform", False),
                            "transformer": p.get("transformer", ""),
                        }
                        for p in rule.get("properties", [])
                    ],
                    "events": [
                        {
                            "vcl": e.get("vcl", ""),
                            "fmx": e.get("fmx", ""),
                            "signature_match": e.get("signature_match", True),
                        }
                        for e in rule.get("events", [])
                    ],
                }
            )
    return rows


def contract_counts() -> tuple[int, Counter]:
    files = sorted(CONTRACTS.rglob("*.expected.json"))
    counts = Counter()
    for path in files:
        try:
            rel = path.relative_to(CONTRACTS)
            counts[rel.parts[0]] += 1
        except Exception:
            pass
    return len(files), counts


def setup_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Segoe UI"
    styles["Normal"].font.size = Pt(10)
    for name, size, color in [
        ("Title", 24, RGBColor(31, 78, 121)),
        ("Heading 1", 17, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(47, 84, 150)),
        ("Heading 3", 11, RGBColor(89, 89, 89)),
    ]:
        styles[name].font.name = "Segoe UI"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = color
    return doc


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Segoe UI"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(31, 78, 121)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(11)
    doc.add_paragraph("Version 5.0 Vanguard", style="Heading 2").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Revision date: {TODAY}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def build_runtime_flow() -> Path:
    builtins = parse_builtin_mappings()
    packs = parse_mapping_packs()
    total_contracts, by_area = contract_counts()
    doc = setup_doc()
    add_title(
        doc,
        "VCL2FMXConverter v5.0 Runtime Flow",
        "Operational flow from startup through contract-backed conversion, mapping, report generation, and validation",
    )

    doc.add_heading("Foreword", level=1)
    doc.add_paragraph(
        "This document updates the v4.1.8 runtime-flow reference for the v5.0 Vanguard line. It explains how the converter starts, loads mapping knowledge, scans a project, converts Pascal and DFM content, applies v5.0 include and Windows-message analysis, generates output, and reports the result."
    )
    doc.add_paragraph(
        "Version 5.0 is contract driven. The converter still uses the same runtime pipeline for normal user conversions, but structural behavior is now guarded by executable conversion contracts before release."
    )

    doc.add_heading("v5.0 Current Standing", level=1)
    rows = [
        ["Release line", "Version 5.0 Vanguard"],
        ["Contract expectations", str(total_contracts)],
        ["Mapping packs", f"{len(list(MAPPING_PACKS.glob('*_MappingPack_v4_1.json')))} JSON packs carried forward from v4.1.8"],
        ["Built-in component mappings", str(len(builtins))],
        ["Mapping-pack component rules", str(len(packs))],
        ["Primary v5.0 additions", "Contracts, include analysis, Windows messaging categories, protected uses cleanup, GDI reporting/conversion contracts, and updated UI/reference surfaces."],
    ]
    add_table(doc, ["Area", "Current v5.0 state"], rows, [2.2, 4.7])

    doc.add_heading("System at a Glance", level=1)
    add_table(
        doc,
        ["Phase", "Runtime purpose", "v5.0 notes"],
        [
            ["Startup", "Create the main FMX form, initialize tabs, options, UI state, and reference browsers.", "The Vanguard UI exposes Dashboard, Project Scan, Component Map, Property Map, Event Map, Conversion Output, and Rules."],
            ["Input selection", "Collect source folder, output folder, file scope, subfolder option, and rule toggles.", "Normal conversion remains user-project based; contracts are not runtime data for user conversions."],
            ["Mapper load", "Load built-in component mappings and optional JSON mapping packs.", "22 mapping packs remain available; pack usage is reported."],
            ["Pascal analysis", "Parse and rewrite source units, uses clauses, Winapi/VCL references, runtime helper needs, and report issues.", "v5.0 analyzes include directives, nested includes, recursive includes, and include boundaries."],
            ["DFM/FMXL generation", "Convert DFM form trees into FMX form output.", "TMemo/TStrings collection handling, TStringGrid event ordering, and UTF-8 BOM FMX output remain preserved."],
            ["Specialized rewrites", "Apply data-aware, LiveBindings, compatibility, WinAPI, GDI, and runtime-normalization rules.", "Windows messaging remains category based, with false-positive protections and explicit manual-review reporting."],
            ["Project generation", "Create the FMX project shell and copy safe companion files.", "Source-local include files and source-root companions are copied without searching build-output folders for ambiguous duplicates."],
            ["Reporting", "Write text/HTML reports and update the UI.", "Reports include blocking status, manual-review status, issue details, mapping-pack usage, Windows-message categories, include findings, and GDI review notes."],
        ],
        [1.4, 2.5, 3.0],
    )

    doc.add_heading("Startup Sequence", level=1)
    for text in [
        "The application starts by creating the main form and initializing the tabbed workspace. Form creation wires runtime events, prepares the source/target controls, initializes scope options, and loads live reference pages.",
        "The UI reference tabs are not static screenshots. Component Map, Property Map, and Event Map are populated from the current mapper knowledge so they match the converter rule surface more closely than old manual tables.",
        "The Rules tab exposes the four major specialized rule families used during a normal conversion pass. Those switches affect runtime conversion behavior; contract files do not.",
    ]:
        doc.add_paragraph(text)

    doc.add_heading("Mapping Load and Reference Surface", level=1)
    doc.add_paragraph(
        "The mapper loads built-in VCL-to-FMX class mappings first, then mapping packs when enabled. Mapping packs can convert, partially convert, detect only, preserve, or route third-party controls to manual review."
    )
    add_table(
        doc,
        ["Mapping source", "What it contributes"],
        [
            ["Built-in mapper", "Standard VCL class mappings, property rows, event rows, explicit substitutions, and fallback research behavior."],
            ["JSON mapping packs", "Third-party component identification, action policy, confidence, property/event rows, vendor notes, and manual-review reasons."],
            ["Live reference tabs", "Current Component Map, Property Map, and Event Map views derived from mapper data at runtime."],
        ],
        [2.0, 4.9],
    )

    doc.add_heading("v5.0 Include Analysis Flow", level=1)
    for text in [
        "Pascal include directives are analyzed before conversion decisions are trusted. Supported forms include {$I FileName.inc}, {$INCLUDE FileName.inc}, (*$I FileName.inc*), and (*$INCLUDE FileName.inc*).",
        "Include resolution is source-tree aware. Beside-source includes, source-subfolder includes, nested includes, missing includes, outside-tree includes, recursive include loops, conditional include locations, and UTF-8 include content are covered by contracts.",
        "The first v5.0 implementation is analysis-first. Include contents can influence detection and reporting, while original include directives remain in generated Pascal output. Include files are copied to the matching output location when safe.",
    ]:
        doc.add_paragraph(text)

    doc.add_heading("v5.0 Windows Messaging Flow", level=1)
    doc.add_paragraph(
        "Windows messaging detection is category based rather than a handwritten list of every constant. The converter looks for active SendMessage, PostMessage, DispatchMessage, PeekMessage, GetMessage, Control.Perform, WM/CM/CN/common-control message families, TWM/TCM records, message declarations, WndProc overrides, WM_USER offsets, and system-command handlers."
    )
    add_table(
        doc,
        ["Category", "Runtime result"],
        [
            ["Safe conversion", "Only applies where an FMX replacement is reliable and covered by a contract."],
            ["FMX helper replacement", "Used for simple control-message families when the generated helper logic is safer than preserving Winapi calls."],
            ["System.Messaging / bridge", "Reserved for message-bridge cases that need an FMX event/message model."],
            ["Platform-specific preservation", "Kept only when active generated code still requires the Windows unit and the report explains why."],
            ["Manual review", "Default for WndProc, message handlers, WM_USER, system command side effects, message pumps, and ambiguous behavior."],
            ["False positive", "Comments, strings, user-defined methods, and unrelated identifiers are ignored under contract."],
        ],
        [2.1, 4.8],
    )

    doc.add_heading("Uses-Clause and Runtime Cleanup", level=1)
    doc.add_paragraph(
        "v5.0 strengthens protected and conditional uses cleanup. The former leftover Vcl.Themes case is covered by contract expectations, along with conditional VCL/Winapi units and no-op false-positive paths."
    )
    doc.add_paragraph(
        "The converter should add Winapi.Messages, Winapi.Windows, System.Messaging, or FMXMessageBridge only when generated active code still requires those units. It should not inject Windows units for false positives."
    )

    doc.add_heading("DFM/FMXL and Encoding Flow", level=1)
    doc.add_paragraph(
        "The DFM parser continues to preserve multi-line string collections such as Lines.Strings, Items.Strings, SQL.Strings, and Params.Strings. TStringGrid event properties are emitted before generated column child objects so Delphi can open the FMX form."
    )
    doc.add_paragraph(
        "Generated FMX files are written as UTF-8 with BOM so accented French text and other international characters load correctly in Delphi."
    )

    doc.add_heading("Contract Runner Relationship", level=1)
    doc.add_paragraph(
        "A conversion contract consists of a small sample Delphi project, form, or source file together with an expectation file that defines the converter output, report items, uses-clause behavior, and required validation behavior."
    )
    doc.add_paragraph(
        "The contract runner is the test utility that loads each contract, runs the converter against its sample input, and verifies that the actual result matches the expectation file. It does not run during normal user conversions."
    )
    add_table(doc, ["Contract area", "Expectation count"], [[k, str(v)] for k, v in sorted(by_area.items())], [3.0, 1.3])

    doc.add_heading("End-to-End Runtime Flow", level=1)
    steps = [
        "User selects a source folder and output folder.",
        "Converter initializes context, options, loaded mapping packs, issue lists, and output paths.",
        "Built-in and third-party mapping rules are loaded.",
        "Source files are discovered while build-output folders are avoided for ambiguous companion-file discovery.",
        "Pascal files are analyzed, include directives are resolved, and rewrite rules are applied.",
        "DFM files are parsed and emitted as FMX with string-collection and event-order protections.",
        "Project files, companion include files, and safe source-root companions are copied or generated.",
        "Reports are written with conversion status, blocking/manual-review details, mapping-pack usage, include findings, Windows-message findings, and GDI/visual-review items.",
        "The UI updates the dashboard/log/report actions for the operator.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    out = DOCS / "V5_0_CONVERTER_RUNTIME_FLOW.docx"
    doc.save(out)
    return out


def mapping_stats(builtins: list[dict], packs: list[dict]) -> dict:
    props = [p for m in builtins for p in m["properties"]]
    events = [e for m in builtins for e in m["events"]]
    return {
        "builtins": len(builtins),
        "packs": len(packs),
        "properties": len(props),
        "events": len(events),
        "component_types": Counter(m.get("mapping_type") or "unspecified" for m in builtins),
        "pack_actions": Counter(m.get("action") or "unspecified" for m in packs),
        "property_transforms": Counter((p.get("transformer") or "direct") for p in props),
        "event_signature": Counter("signature compatible" if e.get("signature_match") else "review signature" for e in events),
    }


def build_maps_reference() -> tuple[Path, Path]:
    builtins = parse_builtin_mappings()
    packs = parse_mapping_packs()
    total_contracts, by_area = contract_counts()
    stats = mapping_stats(builtins, packs)
    doc = setup_doc()
    add_title(
        doc,
        "VCL2FMXConverter v5.0 Rules, Component, Property, and Event Maps",
        "Current reference set for converter rules, mapper rows, event mappings, and mapping-pack coverage",
    )
    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "This reference consolidates the v5.0 rules, component map, property map, and event map into one document placed at the docs root. It reflects the current converter source, mapping-pack set, and v5.0 contract additions."
    )
    add_table(
        doc,
        ["Reference surface", "Current count"],
        [
            ["Built-in component mappings", str(stats["builtins"])],
            ["Built-in property rows", str(stats["properties"])],
            ["Built-in event rows", str(stats["events"])],
            ["Mapping-pack rules", str(stats["packs"])],
            ["Contract expectations", str(total_contracts)],
        ],
        [3.0, 1.5],
    )

    doc.add_heading("v5.0 Rule Families", level=1)
    rules = [
        ["Contract-first development", "Structural behavior should have a sample Delphi input, expected output/report behavior, and a regression guard."],
        ["Include analysis", "Analyze and report Pascal include directives, nested includes, recursive loops, outside-tree paths, conditionals, comments, and UTF-8 content while preserving original include directives."],
        ["Windows messaging", "Detect real message APIs and message declarations by category; report unsafe behavior; ignore comments, strings, user methods, and unrelated identifiers."],
        ["Uses cleanup", "Remove unused VCL/Winapi units, including protected/conditional cases such as Vcl.Themes; add Windows/System.Messaging units only when generated active code requires them."],
        ["GDI drawing", "Convert reliable GDI drawing patterns to FMX canvas calls where covered; still report visual review for drawing behavior that needs human verification."],
        ["DFM/FMXL generation", "Preserve collection-style string properties, emit TStringGrid events before generated columns, and write FMX as UTF-8 with BOM."],
        ["Companion files", "Copy source-local include files and source-root runtime companions without searching build-output folders for ambiguous duplicates."],
        ["Mapping packs", "Load 22 JSON packs, validate actions, report loaded packs and usage, and keep third-party handling visible rather than hidden."],
    ]
    add_table(doc, ["Rule family", "v5.0 behavior"], rules, [2.1, 4.8])

    doc.add_heading("Contract Areas", level=1)
    add_table(doc, ["Area", "Expectation count"], [[k, str(v)] for k, v in sorted(by_area.items())], [3.0, 1.3])

    doc.add_heading("Component Map Summary", level=1)
    add_table(doc, ["Mapping type", "Count"], [[k, str(v)] for k, v in sorted(stats["component_types"].items())], [2.3, 1.0])
    component_rows = []
    for m in builtins:
        component_rows.append([
            m["vcl_class"],
            m["fmx_class"] or "Manual review / none",
            m["mapping_type"] or "",
            str(m["confidence"]),
            m["notes"] or "",
        ])
    add_table(doc, ["VCL class", "FMX target", "Mapping", "Confidence", "Notes"], component_rows[:160], [1.4, 1.4, 1.0, 0.7, 2.6])

    doc.add_heading("Property Map", level=1)
    prop_rows = []
    for m in builtins:
        for p in m["properties"]:
            prop_rows.append([
                m["vcl_class"],
                p.get("vcl", ""),
                p.get("fmx", "") or "Manual review / omitted",
                "yes" if p.get("transform") else "no",
                p.get("transformer", ""),
            ])
    add_table(doc, ["Component", "VCL property", "FMX property", "Transform", "Transformer"], prop_rows[:220], [1.4, 1.4, 1.5, 0.8, 1.5])

    doc.add_heading("Event Map", level=1)
    event_rows = []
    for m in builtins:
        for e in m["events"]:
            event_rows.append([
                m["vcl_class"],
                e.get("vcl", ""),
                e.get("fmx", "") or "Manual review / none",
                "yes" if e.get("signature_match") else "review",
            ])
    add_table(doc, ["Component", "VCL event", "FMX event", "Signature"], event_rows[:220], [1.7, 1.6, 1.6, 1.0])

    doc.add_heading("Mapping-Pack Summary", level=1)
    add_table(doc, ["Action", "Count"], [[k, str(v)] for k, v in sorted(stats["pack_actions"].items())], [2.3, 1.0])
    pack_rows = []
    for m in packs:
        pack_rows.append([
            m.get("pack", ""),
            m.get("vcl_class", ""),
            m.get("fmx_class", "") or "Manual review / none",
            m.get("action", ""),
            str(m.get("confidence", "")),
            m.get("manual_review_reason") or m.get("notes", ""),
        ])
    add_table(doc, ["Pack", "VCL class", "FMX target", "Action", "Confidence", "Notes"], pack_rows[:260], [1.4, 1.3, 1.4, 0.8, 0.7, 2.0])

    doc.add_heading("Notes on Scope", level=1)
    doc.add_paragraph(
        "This document is a reference surface, not a promise that every listed conversion is final production behavior. Direct mappings still require compile and runtime validation, substitute mappings need inspection, and manual-review rows are intentionally conservative."
    )
    doc.add_paragraph(
        "The live converter UI remains the best quick view during operation because it is generated from the running mapper. This document gives a release-reference copy in the docs folder."
    )

    out_docx = DOCS / "VCL2FMXConverter_v5_0_Rules_Component_Property_Event_Maps.docx"
    doc.save(out_docx)

    out_html = DOCS / "VCL2FMXConverter_v5_0_Rules_Component_Property_Event_Maps.html"
    write_maps_html(out_html, builtins, packs, stats, total_contracts, by_area)
    return out_docx, out_html


def html_table(headers: list[str], rows: list[list[str]]) -> str:
    out = ["<table>", "<thead><tr>"]
    out += [f"<th>{html.escape(h)}</th>" for h in headers]
    out += ["</tr></thead><tbody>"]
    for row in rows:
        out.append("<tr>" + "".join(f"<td>{html.escape(str(c))}</td>" for c in row) + "</tr>")
    out += ["</tbody></table>"]
    return "\n".join(out)


def write_maps_html(path: Path, builtins: list[dict], packs: list[dict], stats: dict, total_contracts: int, by_area: Counter) -> None:
    component_rows = [[m["vcl_class"], m["fmx_class"] or "Manual review / none", m["mapping_type"], m["confidence"], m["notes"]] for m in builtins]
    prop_rows = [[m["vcl_class"], p.get("vcl", ""), p.get("fmx", "") or "Manual review / omitted", "yes" if p.get("transform") else "no", p.get("transformer", "")] for m in builtins for p in m["properties"]]
    event_rows = [[m["vcl_class"], e.get("vcl", ""), e.get("fmx", "") or "Manual review / none", "yes" if e.get("signature_match") else "review"] for m in builtins for e in m["events"]]
    pack_rows = [[m.get("pack", ""), m.get("vcl_class", ""), m.get("fmx_class", "") or "Manual review / none", m.get("action", ""), m.get("confidence", ""), m.get("manual_review_reason") or m.get("notes", "")] for m in packs]
    css = """
body{font-family:Segoe UI,Arial,sans-serif;margin:34px;color:#24364b;background:#fff}
h1{color:#1f4e79} h2{color:#2f5496;margin-top:32px}
.meta{display:flex;gap:14px;flex-wrap:wrap;margin:18px 0}
.card{border:1px solid #c9d7ea;border-radius:8px;padding:12px 16px;background:#f7fbff}
table{border-collapse:collapse;width:100%;margin:14px 0 24px 0;font-size:13px}
th{background:#1f4e79;color:white;text-align:left;padding:7px}
td{border:1px solid #d6dce6;padding:6px;vertical-align:top}
tr:nth-child(even){background:#f8fbff}
"""
    body = [
        "<!doctype html><html><head><meta charset='utf-8'>",
        "<title>VCL2FMXConverter v5.0 Rules and Maps</title>",
        f"<style>{css}</style></head><body>",
        "<h1>VCL2FMXConverter v5.0 Rules, Component, Property, and Event Maps</h1>",
        f"<p><strong>Version:</strong> 5.0 Vanguard<br><strong>Revision date:</strong> {html.escape(TODAY)}</p>",
        "<div class='meta'>",
        f"<div class='card'><strong>Built-in components</strong><br>{stats['builtins']}</div>",
        f"<div class='card'><strong>Built-in properties</strong><br>{stats['properties']}</div>",
        f"<div class='card'><strong>Built-in events</strong><br>{stats['events']}</div>",
        f"<div class='card'><strong>Mapping-pack rules</strong><br>{stats['packs']}</div>",
        f"<div class='card'><strong>Contracts</strong><br>{total_contracts}</div>",
        "</div>",
        "<h2>v5.0 Rule Families</h2>",
        html_table(["Rule family", "v5.0 behavior"], [
            ["Contract-first development", "Structural behavior should have a sample Delphi input, an expectation file, and a regression guard."],
            ["Include analysis", "Analyze Pascal include directives, nested includes, recursion, outside-tree paths, conditionals, comments, and UTF-8 content while preserving include directives."],
            ["Windows messaging", "Detect real message APIs and declarations by category, report unsafe behavior, and ignore false positives."],
            ["Uses cleanup", "Remove unused VCL/Winapi units, including protected/conditional Vcl.Themes cases."],
            ["GDI drawing", "Convert reliable GDI patterns to FMX canvas calls where covered and report visual review where needed."],
            ["DFM/FMXL generation", "Preserve string collections, TStringGrid event order, and UTF-8 BOM output."],
            ["Mapping packs", "Load 22 JSON packs, validate actions, and report loaded packs and usage."],
        ]),
        "<h2>Contract Areas</h2>",
        html_table(["Area", "Expectation count"], [[k, v] for k, v in sorted(by_area.items())]),
        "<h2>Component Map</h2>",
        html_table(["VCL class", "FMX target", "Mapping", "Confidence", "Notes"], component_rows),
        "<h2>Property Map</h2>",
        html_table(["Component", "VCL property", "FMX property", "Transform", "Transformer"], prop_rows),
        "<h2>Event Map</h2>",
        html_table(["Component", "VCL event", "FMX event", "Signature"], event_rows),
        "<h2>Mapping-Pack Component Rules</h2>",
        html_table(["Pack", "VCL class", "FMX target", "Action", "Confidence", "Notes"], pack_rows),
        "</body></html>",
    ]
    path.write_text("\n".join(body), encoding="utf-8")


def export_pdf(docx_path: Path) -> Path:
    pdf_path = docx_path.with_suffix(".pdf")
    script = ROOT / "tools" / "export_one_doc_pdf_no_toc.ps1"
    subprocess.run(
        [
            "powershell.exe",
            "-NoProfile",
            "-ExecutionPolicy",
            "Bypass",
            "-File",
            str(script),
            "-InputPath",
            str(docx_path),
            "-PdfPath",
            str(pdf_path),
        ],
        check=True,
        cwd=str(ROOT),
    )
    return pdf_path


def main() -> None:
    runtime = build_runtime_flow()
    maps_docx, maps_html = build_maps_reference()
    print(runtime)
    print(maps_docx)
    print(maps_html)


if __name__ == "__main__":
    main()
